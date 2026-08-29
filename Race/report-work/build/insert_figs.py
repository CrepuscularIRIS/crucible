#!/usr/bin/env python3
"""向已填充的 docx 插入图片,并把最终 docx 导出为 HTML(Chrome 打印 PDF 用)。"""
from __future__ import annotations

import html
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.table import Table
from docx.text.paragraph import Paragraph

DOCX = "/home/lingxufeng/crucible/Race/report-work/build/XH-202619-1B-技术报告.docx"
FIG = "/home/lingxufeng/crucible/Race/report-work/figures"

INSERTS = [  # (docx锚点段前缀, 图片, 宽cm, 说明)
    ("本作品形成的完整闭环", f"{FIG}/fig0-loop-strip.png", 16.5, "图0 总体思路:journal 驱动的两轮闭环(红=反馈返回边)"),
    ("本作品真实架构", f"{FIG}/fig1-architecture.png", 16.5, "图1 五层架构与四条返回边(闭合来自不可绕过的边界)"),
    ("上下文组织", f"{FIG}/fig2-context-layers.png", 13.0, "图2 真实一回合的六层上下文注入结构(journal replay→锚)"),
    ("P15｜", "/home/lingxufeng/crucible/Race/report-work/shots/demo2r/demo2r-16-r2-declared-green.png", 14.0,
     "图3 两轮闭环 UI 实况:第二轮 declare 一次过门(第一轮两拒后修复)"),
    ("P18｜", f"{FIG}/fig6-two-rounds.png", 15.5, "图4 两轮闭环对照(journal 真值,脚本可复算)"),
    ("P18｜", f"{FIG}/fig7-interventions.png", 16.0, "图5 干预前后确定性计数(declare 收口 9/15→6/6;E1 催促缺口如实标注)"),
    ("P19｜", f"{FIG}/fig3-mse-12arm.png", 14.5, "图6 12 臂六世界 MSE 对照(模型×栈两列)"),
    ("P20｜", "/home/lingxufeng/crucible/docs/imgs/image.png", 14.0,
     "图7 百炼调用凭证:6,143 次调用/3.96 亿 token/7 模型全 Qwen 族(08-21→27)"),
]

# ── 1) docx 插图 ───────────────────────────────────────────────
doc = Document(DOCX)
body = doc.element.body
used_anchor = set()
for prefix, path, width_cm, caption in INSERTS:
    for el in list(body):
        if el.tag != qn("w:p") or (prefix, el) in used_anchor:
            continue
        p = Paragraph(el, doc)
        if not p.text.strip().startswith(prefix):
            continue
        used_anchor.add((prefix, el))
        pic_p = doc.add_paragraph()
        el.addnext(pic_p._element)
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(path, width=Cm(width_cm))
        cap_p = doc.add_paragraph()
        pic_p._element.addnext(cap_p._element)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.add_run(caption)
        break
doc.save(DOCX)
print("docx: figures inserted")

# ── 2) docx → HTML(同源导出) ───────────────────────────────────
doc = Document(DOCX)
parts = ['''<!doctype html><html><head><meta charset="utf-8"><style>
@page { size: A4; margin: 14mm 13mm 14mm 13mm; }
body { font-family: "Noto Sans CJK SC","Noto Sans CJK JP","AR PL SungtiL GB",serif; font-size: 10pt; line-height: 1.36; color: #1f2937; margin:0; }
h1.sec { font-size: 13pt; color: #1d4ed8; border-bottom: 2px solid #1d4ed8; padding-bottom: 2px; margin: 0 0 7px 0; }
div.page { page-break-before: always; }
table { border-collapse: collapse; width: 100%; margin: 5px 0; }
td, th { border: 1px solid #94a3b8; padding: 2.5px 4.5px; font-size: 8.4pt; vertical-align: top; }
th { background: #eff6ff; font-weight: 600; }
p { margin: 3.5px 0; }
img { max-width: 100%; }
.cap { text-align: center; font-size: 8.3pt; color: #475569; margin: 1px 0 5px 0; }
</style></head><body>''']

fig_by_caption = {cap: (path, w) for _, path, w, cap in INSERTS}
captions_emitted = set()
for el in doc.element.body:
    tag = el.tag.split("}")[1]
    if tag == "p":
        p = Paragraph(el, doc)
        t = p.text.strip()
        has_pic = bool(p._element.findall(".//" + qn("a:blip")))
        if has_pic:  # 图片段:找 INSERTS 里未发出的第一张
            for _, path, w, cap in INSERTS:
                if cap not in captions_emitted:
                    parts.append(f'<p><img src="file://{path}" style="width:{w}cm"></p>')
                    parts.append(f'<p class="cap">{html.escape(cap)}</p>')
                    captions_emitted.add(cap)
                    break
            continue
        if not t:
            continue
        if t in fig_by_caption:  # docx caption 段跳过(HTML 已带)
            continue
        if t.startswith("P") and "｜" in t[:8]:
            parts.append(f'<div class="page"><h1 class="sec">{html.escape(t)}</h1>')
        elif t.startswith("【图"):  # 正文里的图片占位说明
            parts.append(f"<p>{html.escape(t)}</p>")
        else:
            parts.append(f"<p>{html.escape(t)}</p>")
    elif tag == "tbl":
        tb = Table(el, doc)
        rows_html = []
        for r_i, row in enumerate(tb.rows):
            cells = []
            for c in row.cells:
                txt = html.escape(c.text.strip())
                cells.append(f"<th>{txt}</th>" if r_i == 0 else f"<td>{txt}</td>")
            rows_html.append("<tr>" + "".join(cells) + "</tr>")
        parts.append("<table>" + "".join(rows_html) + "</table>")
parts.append("</body></html>")
out = DOCX.replace(".docx", ".html")
open(out, "w").write("\n".join(parts))
print("html written:", out, "| figs emitted:", len(captions_emitted))
